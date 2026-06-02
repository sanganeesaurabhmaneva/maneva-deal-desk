#!/usr/bin/env python3
"""build_proposal.py -- appends the proposal under Appendix 1 of an already-filled
Maneva Service Agreement, matching the XConnect output layout (navy-header tables,
Executive Summary, one block per application). The narrative comes from the rep's typed
input; the Estimated Pricing comes from the calculator. Photos, if any, go under their
labelled section.

Usage (standalone): python3 build_proposal.py <agreement.docx> <payload.json> <out.docx>
Or import build_into(doc_path, payload, out_path).
"""
import sys, json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1A, 0x3C, 0x5E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER = "CCCCCC"


def _set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), BORDER)
        borders.append(e)
    tblPr.append(borders)


def _cell_text(cell, text, bold=False, color=None, size=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("" if text is None else str(text))
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)


def _heading(doc, text, level):
    sizes = {1: 15, 2: 13, 3: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = NAVY
    return p


def _para(doc, text, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("" if text is None else str(text))
    r.font.size = Pt(size)
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def _bullets(doc, items):
    for it in (items or []):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("\u2022  " + str(it)).font.size = Pt(10.5)


def _table(doc, headers, rows, widths=None):
    cols = len(headers)
    t = doc.add_table(rows=1, cols=cols)
    _set_table_borders(t)
    t.allow_autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_bg(hdr[i], "1A3C5E")
        _cell_text(hdr[i], h, bold=True, color=WHITE, size=10)
    for row in rows:
        cells = t.add_row().cells
        for i in range(cols):
            val = row[i] if i < len(row) else ""
            _cell_text(cells[i], val, size=10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    return t


def _money(n, currency):
    sym = "CA$" if str(currency).upper() == "CAD" else "$"
    try:
        return f"{sym}{round(float(n)):,}"
    except (TypeError, ValueError):
        return str(n)


def _pricing_rows(pricing, currency):
    """Estimated Pricing table (Component / Details / Cost) from the calculator."""
    rate = pricing.get("finalRatePerLine") or pricing.get("final_rate_per_line") or 0
    lines = pricing.get("lines") or 1
    mrr = pricing.get("mrr") or rate * lines
    arr = pricing.get("arr") or mrr * 12
    rows = [
        ["Software License (Monthly)", f"{lines} line(s) at {_money(rate, currency)}/line/month", f"{_money(mrr, currency)}/month"],
        ["Software License (Annual)", f"{_money(mrr, currency)} x 12 months", f"{_money(arr, currency)}/year"],
        ["Installation", "Included", "Included"],
    ]
    return rows


def _impact_rows(impact, pricing, currency, hardware_total):
    """Key ROI Metrics table (label / value). Prefer rep-provided; compute where possible."""
    arr = pricing.get("arr") or 0
    rows = []
    tav = impact.get("total_annual_value")
    if tav is not None:
        rows.append(["Total Annual Value", _money(tav, currency)])
    y1 = impact.get("year1_investment")
    if y1 is None and (hardware_total or arr):
        y1 = (hardware_total or 0) + arr
    if y1 is not None:
        rows.append(["Year 1 Investment (Hardware + Subscription)", _money(y1, currency)])
    net = impact.get("annual_net_benefit")
    if net is None and tav is not None:
        try:
            net = float(tav) - float(arr)
        except (TypeError, ValueError):
            net = None
    if net is not None:
        rows.append(["Annual Net Benefit", _money(net, currency)])
    payback = impact.get("payback_period")
    if payback:
        rows.append(["Payback Period", payback])
    elif tav:
        try:
            months = (float(y1) / (float(tav) / 12.0)) if tav else None
            if months:
                rows.append(["Payback Period", f"{months:.1f} months"])
        except (TypeError, ValueError, ZeroDivisionError):
            pass
    return rows


def build_into(doc_path, payload, out_path):
    proposal = (payload or {}).get("proposal") or {}
    pricing = (payload or {}).get("pricing") or {}
    currency = pricing.get("currency") or (payload or {}).get("currency") or "USD"
    photos = proposal.get("photos") or []

    doc = Document(doc_path)

    # start the proposal on a new page
    doc.add_page_break()

    if proposal.get("executive_summary"):
        _heading(doc, "Executive Summary", 1)
        es = proposal["executive_summary"]
        for para in (es if isinstance(es, list) else [es]):
            _para(doc, para)

    _heading(doc, "Project Proposal", 2)

    for app in (proposal.get("applications") or []):
        _heading(doc, app.get("name", "Application"), 1)

        if app.get("problem"):
            _heading(doc, "Problem Statement", 2); _para(doc, app["problem"])
        if app.get("solution") or app.get("inspection_points"):
            _heading(doc, "Proposed Solution", 2)
            if app.get("solution"):
                _para(doc, app["solution"])
            if app.get("inspection_points"):
                _para(doc, "The AI model will inspect each unit for:")
                _bullets(doc, app["inspection_points"])
        if app.get("before_after"):
            _heading(doc, "Process Comparison: Before vs. After", 2)
            _table(doc, ["BEFORE (Current Process)", "AFTER (With Maneva AI)"],
                   [[r[0], r[1]] for r in app["before_after"]], widths=[3.25, 3.25])
        if app.get("hardware"):
            _heading(doc, "Hardware Requirements", 2)
            _table(doc, ["Item", "Qty", "Description"],
                   [[r[0], r[1] if len(r) > 1 else "", r[2] if len(r) > 2 else ""] for r in app["hardware"]],
                   widths=[1.6, 0.8, 4.1])
        # Estimated Pricing always from the calculator
        _heading(doc, "Estimated Pricing", 2)
        _table(doc, ["Component", "Details", "Cost"],
               app.get("pricing_rows") or _pricing_rows(pricing, currency), widths=[2.1, 2.8, 1.6])

        _heading(doc, "Impact & Benefits", 3)
        hw_total = 0
        for r in (app.get("hardware") or []):
            pass  # hardware costs live in the agreement's hardware table; impact uses rep figures
        impact_rows = _impact_rows(app.get("impact") or {}, pricing, currency, app.get("hardware_total") or 0)
        if impact_rows:
            _table(doc, ["Key ROI Metrics", ""], impact_rows, widths=[3.5, 3.0])
        if app.get("preventative_benefits"):
            _heading(doc, "Preventative Benefits \u2014 Cost Avoidance", 3)
            pb = app["preventative_benefits"]
            _bullets(doc, pb) if isinstance(pb, list) else _para(doc, pb)
        if app.get("upside_benefits"):
            _heading(doc, "Upside Benefits \u2014 Value Creation", 3)
            ub = app["upside_benefits"]
            _bullets(doc, ub) if isinstance(ub, list) else _para(doc, ub)
        if app.get("assumptions"):
            _heading(doc, "Key Assumptions & Data Sources", 3)
            asm = app["assumptions"]
            _bullets(doc, asm) if isinstance(asm, list) else _para(doc, asm)
        if app.get("timeline"):
            _heading(doc, "Implementation Timeline", 2)
            tl = app["timeline"]
            _bullets(doc, tl) if isinstance(tl, list) else _para(doc, tl)
        if app.get("risks"):
            _heading(doc, "Risks & Mitigations", 2)
            _table(doc, ["Risk", "Mitigation"],
                   [[r[0], r[1] if len(r) > 1 else ""] for r in app["risks"]], widths=[3.25, 3.25])

        # photos for this application, placed under a labelled subsection
        app_photos = [p for p in photos if p.get("application", app.get("name")) == app.get("name")] or \
                     (photos if len(proposal.get("applications") or []) == 1 else [])
        if app_photos:
            _heading(doc, "Supporting Photos", 3)
            for ph in app_photos:
                path = ph.get("path")
                if path and os.path.exists(path):
                    try:
                        pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pic_p.add_run().add_picture(path, width=Inches(5.0))
                        if ph.get("label"):
                            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cr = cap.add_run(ph["label"]); cr.italic = True; cr.font.size = Pt(9)
                            cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    except Exception as e:
                        _para(doc, f"[photo could not be embedded: {ph.get('label','')}]", italic=True)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("usage: build_proposal.py <agreement.docx> <payload.json> <out.docx>"); sys.exit(1)
    payload = json.load(open(sys.argv[2], encoding="utf-8"))
    build_into(sys.argv[1], payload, sys.argv[3])
    print("wrote", sys.argv[3])
